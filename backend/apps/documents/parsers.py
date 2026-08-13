import hashlib
import io
import json
import zipfile
from collections.abc import Callable
from dataclasses import dataclass
from typing import BinaryIO
from xml.etree import ElementTree

from django.conf import settings

from .ocr import OcrProvider
from .parse_exceptions import (
    DocumentParseContentInvalid,
    DocumentParseSecurityRejected,
)

PARSER_VERSION = "1"
_ALLOWED_CONTROLS = {"\n", "\t"}
_DANGEROUS_PDF_KEYS = {"/JS", "/JavaScript", "/Launch", "/EmbeddedFiles", "/OpenAction", "/AA"}


@dataclass(frozen=True)
class ParseResult:
    canonical_text: str
    tables: list[list[list[str]]]
    warning_codes: list[str]
    parser_key: str
    parser_version: str
    ocr_provider_key: str = ""
    ocr_engine_version: str = ""


def canonicalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n")
    if "\x00" in value:
        raise DocumentParseContentInvalid
    if any(ord(char) < 32 and char not in _ALLOWED_CONTROLS for char in value):
        raise DocumentParseContentInvalid
    if len(value) > settings.DOCUMENT_PARSE_MAX_CHARACTERS:
        raise DocumentParseContentInvalid
    try:
        encoded = value.encode("utf-8")
    except UnicodeEncodeError as exc:
        raise DocumentParseContentInvalid from exc
    if len(encoded) > settings.DOCUMENT_PARSE_MAX_UTF8_BYTES:
        raise DocumentParseContentInvalid
    return value


def _cell(value) -> str:
    text = "" if value is None else canonicalize_text(str(value))
    if len(text) > settings.DOCUMENT_PARSE_MAX_CELL_CHARACTERS:
        raise DocumentParseContentInvalid
    return text


def _bounded_tables(tables) -> list[list[list[str]]]:
    result: list[list[list[str]]] = []
    for table in tables:
        if len(result) >= settings.DOCUMENT_PARSE_MAX_TABLES:
            raise DocumentParseContentInvalid
        rows: list[list[str]] = []
        for row in table:
            if len(rows) >= settings.DOCUMENT_PARSE_MAX_TABLE_ROWS:
                raise DocumentParseContentInvalid
            cells = [_cell(value) for value in row]
            if len(cells) > settings.DOCUMENT_PARSE_MAX_TABLE_COLUMNS:
                raise DocumentParseContentInvalid
            rows.append(cells)
        result.append(rows)
    serialized = json.dumps(result, ensure_ascii=False, separators=(",", ":")).encode()
    if len(serialized) > settings.DOCUMENT_PARSE_MAX_TABLE_JSON_BYTES:
        raise DocumentParseContentInvalid
    return result


def _safe_zip(data: bytes, *, reject_external_relationships: bool) -> zipfile.ZipFile:
    try:
        archive = zipfile.ZipFile(io.BytesIO(data))
    except (zipfile.BadZipFile, ValueError) as exc:
        raise DocumentParseContentInvalid from exc
    total = 0
    for entry in archive.infolist():
        normalized = entry.filename.replace("\\", "/")
        if normalized.startswith("/") or ".." in normalized.split("/"):
            raise DocumentParseSecurityRejected
        total += entry.file_size
        if (
            entry.file_size > settings.FILE_VALIDATION_MAX_ARCHIVE_ENTRY_BYTES
            or total > settings.FILE_VALIDATION_MAX_UNCOMPRESSED_BYTES
        ):
            raise DocumentParseSecurityRejected
        lowered = normalized.casefold()
        if (
            "vbaproject" in lowered
            or "/embeddings/" in lowered
            or lowered.endswith((".bin", ".exe"))
        ):
            raise DocumentParseSecurityRejected
        if lowered.startswith("xl/externallinks/"):
            raise DocumentParseSecurityRejected
        if lowered.startswith("xl/externallinks/"):
            raise DocumentParseSecurityRejected
        if reject_external_relationships and lowered.endswith(".rels"):
            try:
                root = ElementTree.fromstring(archive.read(entry))
            except ElementTree.ParseError as exc:
                raise DocumentParseContentInvalid from exc
            if any(node.attrib.get("TargetMode") == "External" for node in root):
                raise DocumentParseSecurityRejected
    return archive


def parse_pdf(stream: BinaryIO, _: OcrProvider | None = None) -> ParseResult:
    from pypdf import PdfReader

    try:
        reader = PdfReader(stream, strict=True)
    except Exception as exc:
        raise DocumentParseContentInvalid from exc
    if reader.is_encrypted:
        raise DocumentParseSecurityRejected
    if len(reader.pages) > 1000:
        raise DocumentParseContentInvalid
    raw_structure = repr(reader.trailer)
    if any(key in raw_structure for key in _DANGEROUS_PDF_KEYS):
        raise DocumentParseSecurityRejected
    chunks = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")
    text = canonicalize_text("\n".join(chunks))
    warnings = [] if text else ["NO_TEXT_EXTRACTED", "OCR_NOT_APPLIED"]
    return ParseResult(text, [], warnings, "pdf", PARSER_VERSION)


def parse_docx(stream: BinaryIO, _: OcrProvider | None = None) -> ParseResult:
    from docx import Document

    data = stream.read(settings.FILE_UPLOAD_MAX_BYTES + 1)
    archive = _safe_zip(data, reject_external_relationships=True)
    archive.close()
    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise DocumentParseContentInvalid from exc
    paragraphs = [paragraph.text for paragraph in document.paragraphs]
    tables = [
        [[cell.text for cell in row.cells] for row in table.rows] for table in document.tables
    ]
    return ParseResult(
        canonicalize_text("\n".join(paragraphs)),
        _bounded_tables(tables),
        [],
        "docx",
        PARSER_VERSION,
    )


def parse_xlsx(stream: BinaryIO, _: OcrProvider | None = None) -> ParseResult:
    from openpyxl import load_workbook  # type: ignore[import-untyped]

    data = stream.read(settings.FILE_UPLOAD_MAX_BYTES + 1)
    archive = _safe_zip(data, reject_external_relationships=True)
    archive.close()
    try:
        workbook = load_workbook(io.BytesIO(data), read_only=True, data_only=True, keep_links=False)
    except Exception as exc:
        raise DocumentParseContentInvalid from exc
    if len(workbook.worksheets) > 100:
        raise DocumentParseContentInvalid
    tables = []
    lines = []
    for worksheet in workbook.worksheets:
        rows = []
        for index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
            if index > settings.DOCUMENT_PARSE_MAX_TABLE_ROWS:
                raise DocumentParseContentInvalid
            values = [_cell(value) for value in row]
            if len(values) > settings.DOCUMENT_PARSE_MAX_TABLE_COLUMNS:
                raise DocumentParseContentInvalid
            rows.append(values)
            lines.append("\t".join(values))
        tables.append(rows)
    workbook.close()
    return ParseResult(
        canonicalize_text("\n".join(lines)),
        _bounded_tables(tables),
        [],
        "xlsx",
        PARSER_VERSION,
    )


def parse_plain(stream: BinaryIO, _: OcrProvider | None = None, *, key: str) -> ParseResult:
    data = stream.read(settings.DOCUMENT_PARSE_MAX_UTF8_BYTES + 1)
    if len(data) > settings.DOCUMENT_PARSE_MAX_UTF8_BYTES:
        raise DocumentParseContentInvalid
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise DocumentParseContentInvalid from exc
    return ParseResult(canonicalize_text(text), [], [], key, PARSER_VERSION)


def parse_image(stream: BinaryIO, ocr: OcrProvider | None = None) -> ParseResult:
    if ocr is None or not ocr.is_available():
        from .parse_exceptions import DocumentOcrUnavailable

        raise DocumentOcrUnavailable
    result = ocr.recognize(stream)
    return ParseResult(
        canonicalize_text(result.text),
        [],
        list(result.warning_codes),
        "image_ocr",
        PARSER_VERSION,
        ocr.key,
        result.engine_version,
    )


def parse_txt(stream: BinaryIO, ocr: OcrProvider | None = None) -> ParseResult:
    return parse_plain(stream, ocr, key="txt")


def parse_markdown(stream: BinaryIO, ocr: OcrProvider | None = None) -> ParseResult:
    return parse_plain(stream, ocr, key="markdown")


Parser = Callable[[BinaryIO, OcrProvider | None], ParseResult]
PARSERS: dict[str, Parser] = {
    "pdf": parse_pdf,
    "docx": parse_docx,
    "xlsx": parse_xlsx,
    "txt": parse_txt,
    "markdown": parse_markdown,
    "jpg": parse_image,
    "jpeg": parse_image,
    "png": parse_image,
    "webp": parse_image,
}


def parser_key_for(file_kind: str) -> str:
    try:
        parser = PARSERS[file_kind]
    except KeyError as exc:
        raise DocumentParseContentInvalid from exc
    return "image_ocr" if parser is parse_image else file_kind


def parse_stream(file_kind: str, stream: BinaryIO, ocr: OcrProvider | None) -> ParseResult:
    try:
        parser = PARSERS[file_kind]
    except KeyError as exc:
        raise DocumentParseContentInvalid from exc
    return parser(stream, ocr)


def machine_digest(*, document_version, result: ParseResult) -> str:
    payload = {
        "document_version_id": str(document_version.pk),
        "file_sha256": document_version.sha256,
        "parser_key": result.parser_key,
        "parser_version": result.parser_version,
        "ocr_engine_version": result.ocr_engine_version,
        "text": result.canonical_text,
        "tables": result.tables,
        "warnings": sorted(set(result.warning_codes)),
    }
    return hashlib.sha256(
        json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
    ).hexdigest()


def confirmation_digest(*, machine_base, parent, text: str) -> str:
    payload = {
        "machine_base_id": str(machine_base.pk),
        "machine_base_digest": machine_base.content_digest,
        "parent_id": str(parent.pk),
        "text": text,
    }
    return hashlib.sha256(
        json.dumps(payload, ensure_ascii=False, sort_keys=True, separators=(",", ":")).encode()
    ).hexdigest()
