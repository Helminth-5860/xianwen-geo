from __future__ import annotations

import hashlib
import io
import json
import textwrap
from dataclasses import asdict
from datetime import datetime, timedelta
from decimal import Decimal
from typing import Any, Literal, cast
from uuid import UUID

from django.core.exceptions import ObjectDoesNotExist
from django.db import transaction
from django.db.models import Avg
from django.http import Http404
from django.utils import timezone

from apps.documents.storage import storage_provider

from .competitors import active_competitor_entities, persist_competitor_reference
from .exposure import exposure_for_job
from .models import (
    GeoDetectionJob,
    GeoReport,
    ModelCall,
    ModelScoreResult,
    ReportExport,
    ScoreResult,
)
from .score_aggregation import ModelScore, aggregate_composite_score, score_grade


def _digest(value: Any) -> str:
    encoded = json.dumps(value, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(encoded.encode()).hexdigest()


def _json_value(value: Any) -> Any:
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, datetime):
        return value.isoformat()
    if isinstance(value, UUID):
        return str(value)
    if isinstance(value, dict):
        return {key: _json_value(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_json_value(item) for item in value]
    return value


def question_facts(job: GeoDetectionJob) -> list[dict[str, Any]]:
    return [
        {
            "source_question_id": str(row.source_question_id),
            "text": row.text,
            "question_type": row.question_type,
            "participates_in_scoring": row.participates_in_scoring,
            "primary_category_key": row.primary_category_key,
            "primary_category_name": row.primary_category_name,
            "primary_category_version": row.primary_category_version,
            "priority": row.priority,
            "sort_order": row.sort_order,
        }
        for row in job.snapshot.questions.order_by("sort_order", "id")
    ]


def model_facts(job: GeoDetectionJob) -> list[dict[str, Any]]:
    return [
        {"model_id": str(row.model_id), "model_key": row.model_key}
        for row in job.model_runs.order_by("model_id")
    ]


def _composite(job: GeoDetectionJob, track: Literal["geo", "brand_reputation"]) -> dict[str, Any]:
    rows = ModelScoreResult.objects.filter(model_run__job=job, track=track)
    scores = tuple(
        ModelScore(
            track=cast(Literal["geo", "brand_reputation"], row.track),
            planned_count=row.planned_count,
            successful_count=row.successful_count,
            success_rate=row.success_rate,
            score=row.score,
            status=cast(Literal["formal", "reference", "not_generated"], row.status),
        )
        for row in rows
    )
    result = aggregate_composite_score(track=track, model_scores=scores)
    return {
        "score": str(result.score) if result.score is not None else None,
        "grade": score_grade(result.score) if result.score is not None else None,
        "status": result.status,
        "formal_model_count": result.formal_model_count,
    }


def _dimensions(job: GeoDetectionJob) -> dict[str, str | None]:
    rows = ScoreResult.objects.filter(model_response__model_call__job=job)
    natural_rows = rows.filter(track=ScoreResult.Track.GEO)
    values = {
        "mention": natural_rows.aggregate(value=Avg("mention_score"))["value"],
        "recommendation": rows.aggregate(value=Avg("recommendation_score"))["value"],
        "rank": natural_rows.aggregate(value=Avg("rank_score"))["value"],
        "accuracy": rows.aggregate(value=Avg("accuracy_score"))["value"],
        "sentiment": rows.aggregate(value=Avg("sentiment_score"))["value"],
        "citation": rows.aggregate(value=Avg("citation_score"))["value"],
    }
    return {
        key: str(Decimal(value).quantize(Decimal("0.0001"))) if value is not None else None
        for key, value in values.items()
    }


def _model_summaries(job: GeoDetectionJob) -> list[dict[str, Any]]:
    result = []
    for row in job.model_runs.order_by("model__canonical_order", "id"):
        scores = {
            score.track: {
                "score": str(score.score) if score.score is not None else None,
                "status": score.status,
                "planned_count": score.planned_count,
                "successful_count": score.successful_count,
                "success_rate": (
                    str(score.success_rate) if score.success_rate is not None else None
                ),
            }
            for score in row.score_results.all()
        }
        result.append(
            {
                "model_id": str(row.model_id),
                "model_key": row.model_key,
                "status": row.status,
                "planned_calls": row.planned_calls,
                "completed_calls": row.completed_calls,
                "successful_calls": row.successful_calls,
                "failed_calls": row.failed_calls,
                "cancelled_calls": row.cancelled_calls,
                "geo": scores.get("geo"),
                "brand_reputation": scores.get("brand_reputation"),
            }
        )
    return result


def _semantic_provenance(job: GeoDetectionJob) -> list[dict[str, str]]:
    rows = (
        ScoreResult.objects.filter(model_response__model_call__job=job)
        .values(
            "semantic_schema_version",
            "semantic_provider_key",
            "semantic_model_key",
            "semantic_adapter_version",
            "semantic_prompt_version",
            "semantic_provider_model_id",
        )
        .distinct()
        .order_by(
            "semantic_provider_key",
            "semantic_model_key",
            "semantic_adapter_version",
            "semantic_prompt_version",
            "semantic_provider_model_id",
        )
    )
    return [
        {
            "semantic_schema_version": str(row["semantic_schema_version"]),
            "semantic_provider_key": str(row["semantic_provider_key"]),
            "semantic_model_key": str(row["semantic_model_key"]),
            "semantic_adapter_version": str(row["semantic_adapter_version"]),
            "semantic_prompt_version": str(row["semantic_prompt_version"]),
            "semantic_provider_model_id": str(row["semantic_provider_model_id"]),
        }
        for row in rows
    ]


def report_is_ready(job: GeoDetectionJob) -> bool:
    if job.status not in {
        GeoDetectionJob.Status.PARTIAL,
        GeoDetectionJob.Status.SUCCEEDED,
        GeoDetectionJob.Status.FAILED,
        GeoDetectionJob.Status.CANCELLED,
    }:
        return False
    scoring_calls = job.model_calls.filter(question_snapshot__participates_in_scoring=True)
    if scoring_calls.filter(
        status=ModelCall.Status.SUCCEEDED, response__score_result__isnull=True
    ).exists():
        return False
    for run in job.model_runs.all():
        expected_tracks = set(
            run.calls.filter(question_snapshot__participates_in_scoring=True)
            .values_list("question_snapshot__question_type", flat=True)
            .distinct()
        )
        actual_tracks = set(run.score_results.values_list("track", flat=True))
        required_tracks = {
            "geo" if question_type == "natural" else "brand_reputation"
            for question_type in expected_tracks
        }
        if required_tracks != actual_tracks:
            return False
    return True


@transaction.atomic
def prepare_report(*, job: GeoDetectionJob) -> GeoReport | None:
    job = GeoDetectionJob.objects.select_for_update().get(pk=job.pk)
    existing = GeoReport.objects.filter(job=job).first()
    if existing is not None:
        return existing
    if job.status not in {
        GeoDetectionJob.Status.PARTIAL,
        GeoDetectionJob.Status.SUCCEEDED,
        GeoDetectionJob.Status.FAILED,
        GeoDetectionJob.Status.CANCELLED,
    }:
        return None
    from .score_orchestration import finalize_model_run_scores

    for model_run_id in job.model_runs.values_list("id", flat=True):
        finalize_model_run_scores(model_run_id=model_run_id)
    if not report_is_ready(job):
        return None
    if (
        not job.competitor_entities.exists()
        and ScoreResult.objects.filter(model_response__model_call__job=job).exists()
    ):
        persist_competitor_reference(job=job)
    return generate_report(job=job)


@transaction.atomic
def generate_report(*, job: GeoDetectionJob, baseline: GeoReport | None = None, retest_mode=""):
    if job.status not in {
        GeoDetectionJob.Status.PARTIAL,
        GeoDetectionJob.Status.SUCCEEDED,
        GeoDetectionJob.Status.FAILED,
        GeoDetectionJob.Status.CANCELLED,
    }:
        raise ValueError("report requires a terminal detection")
    if not report_is_ready(job):
        raise ValueError("report scoring is not complete")
    existing = GeoReport.objects.filter(job=job).first()
    if existing is not None:
        return existing
    if baseline is None:
        try:
            origin = job.retest_origin
        except ObjectDoesNotExist:
            origin = None
        if origin is not None:
            baseline = origin.baseline_report
            retest_mode = origin.mode
    questions = question_facts(job)
    models = model_facts(job)
    exposure = exposure_for_job(job=job)
    competitors = [
        {
            "id": str(entity.pk),
            "canonical_name": entity.canonical_name,
            "aliases": entity.aliases,
            "entity_type": entity.entity_type,
            "mention_count": entity.mentions.count(),
        }
        for entity in active_competitor_entities(job=job)
    ]
    summary = {
        "geo": _composite(job, "geo"),
        "brand_reputation": _composite(job, "brand_reputation"),
        "exposure": _json_value(asdict(exposure)),
        "models": _model_summaries(job),
        "dimensions": _dimensions(job),
        "competitors": competitors,
    }
    provenance = {
        "detection_id": str(job.pk),
        "subject_id": str(job.subject_id),
        "subject_version_id": str(job.snapshot.subject_version_id),
        "question_bank_version_id": str(job.snapshot.question_bank_version_id),
        "questions": questions,
        "models": models,
        "prompt_version": job.snapshot.prompt_version,
        "scoring_rule_version": job.snapshot.scoring_rule_version,
        "semantic_scoring": _semantic_provenance(job),
    }
    return GeoReport.objects.create(
        job=job,
        user=job.user,
        subject=job.subject,
        subject_version=job.snapshot.subject_version,
        baseline_report=baseline,
        retest_mode=retest_mode,
        question_signature=_digest(questions),
        model_signature=_digest(models),
        scoring_rule_version=job.snapshot.scoring_rule_version,
        summary=summary,
        provenance=provenance,
    )


def report_for_user_or_404(*, user, report_id) -> GeoReport:
    try:
        return GeoReport.objects.select_related("job", "subject", "subject_version").get(
            pk=report_id, user=user
        )
    except GeoReport.DoesNotExist as exc:
        raise Http404 from exc


def comparison(current: GeoReport, baseline: GeoReport | None = None) -> dict[str, Any] | None:
    baseline = baseline or current.baseline_report
    if baseline is None:
        return None
    same_subject = current.subject_id == baseline.subject_id
    same_questions = current.question_signature == baseline.question_signature
    same_models = current.model_signature == baseline.model_signature
    same_rule = current.scoring_rule_version == baseline.scoring_rule_version
    comparable = same_subject and same_questions and same_models and same_rule

    def numeric_delta(current_value, baseline_value):
        if not comparable or current_value is None or baseline_value is None:
            return None
        return str(Decimal(current_value) - Decimal(baseline_value))

    model_deltas = []
    baseline_models = {row["model_id"]: row for row in baseline.summary["models"]}
    if comparable:
        for row in current.summary["models"]:
            prior = baseline_models[row["model_id"]]
            model_deltas.append(
                {
                    "model_id": row["model_id"],
                    "model_key": row["model_key"],
                    "geo_score_delta": numeric_delta(
                        (row.get("geo") or {}).get("score"),
                        (prior.get("geo") or {}).get("score"),
                    ),
                    "brand_reputation_score_delta": numeric_delta(
                        (row.get("brand_reputation") or {}).get("score"),
                        (prior.get("brand_reputation") or {}).get("score"),
                    ),
                }
            )
    dimension_deltas = {
        key: numeric_delta(value, baseline.summary["dimensions"].get(key))
        for key, value in current.summary["dimensions"].items()
    }
    return {
        "baseline_report_id": str(baseline.pk),
        "status": "comparable" if comparable else "not_comparable",
        "same_subject": same_subject,
        "same_questions": same_questions,
        "same_models": same_models,
        "same_scoring_rule": same_rule,
        "subject_version_changed": current.subject_version_id != baseline.subject_version_id,
        "scoring_version_changed": not same_rule,
        "geo_score_delta": numeric_delta(
            current.summary["geo"]["score"], baseline.summary["geo"]["score"]
        ),
        "brand_reputation_score_delta": numeric_delta(
            current.summary["brand_reputation"]["score"],
            baseline.summary["brand_reputation"]["score"],
        ),
        "exposure_index_delta": numeric_delta(
            current.summary["exposure"]["exposure_index"],
            baseline.summary["exposure"]["exposure_index"],
        ),
        "dimension_deltas": dimension_deltas if comparable else {},
        "model_deltas": model_deltas,
    }


def report_payload(report: GeoReport) -> dict[str, Any]:
    return {
        "id": str(report.pk),
        "detection_id": str(report.job_id),
        "subject_id": str(report.subject_id),
        "subject_version_id": str(report.subject_version_id),
        "retest_mode": report.retest_mode,
        "summary": report.summary,
        "provenance": report.provenance,
        "comparison": comparison(report),
        "generated_at": report.generated_at,
    }


def question_group_page(*, report: GeoReport, page: int, page_size: int = 10) -> dict[str, Any]:
    questions = list(report.job.snapshot.questions.order_by("sort_order", "id"))
    count = len(questions)
    start = (page - 1) * page_size
    groups = []
    for question in questions[start : start + page_size]:
        calls = (
            ModelCall.objects.filter(job=report.job, question_snapshot=question)
            .select_related("model_run", "response__score_result")
            .prefetch_related("response__citations")
        )
        results = []
        for call in calls.order_by("model_run__model__canonical_order", "id"):
            try:
                response = call.response
            except ObjectDoesNotExist:
                response = None
            try:
                score = response.score_result if response is not None else None
            except ObjectDoesNotExist:
                score = None
            results.append(
                {
                    "call_id": str(call.pk),
                    "model_id": str(call.model_id),
                    "model_key": call.model_key,
                    "status": call.status,
                    "safe_error_summary": call.safe_error_summary,
                    "answer_available": response is not None,
                    "snippet": response.raw_text[:300] if response is not None else "",
                    "score": _json_value(
                        {
                            "total": score.total_score,
                            "mention": score.mention_score,
                            "recommendation": score.recommendation_score,
                            "rank": score.rank_score,
                            "accuracy": score.accuracy_score,
                            "sentiment": score.sentiment_score,
                            "citation": score.citation_score,
                        }
                    )
                    if score is not None
                    else None,
                    "citations": [
                        {
                            "title": row.title,
                            "url": row.canonical_url,
                            "source_name": row.source_name,
                            "quoted_text": row.quoted_text,
                        }
                        for row in response.citations.filter(url_status="safe")
                    ]
                    if response is not None
                    else [],
                }
            )
        groups.append(
            {
                "question_id": str(question.pk),
                "source_question_id": str(question.source_question_id),
                "question_type": question.question_type,
                "text": question.text,
                "results": results,
            }
        )
    return {
        "results": groups,
        "pagination": {
            "page": page,
            "page_size": page_size,
            "count": count,
            "total_pages": max(1, (count + page_size - 1) // page_size),
        },
    }


def full_answer(*, report: GeoReport, call_id) -> dict[str, Any]:
    try:
        call = report.job.model_calls.select_related("response").get(pk=call_id)
        response = call.response
    except (ModelCall.DoesNotExist, ObjectDoesNotExist) as exc:
        raise Http404 from exc
    return {
        "call_id": str(call.pk),
        "model_key": call.model_key,
        "answer": response.raw_text,
        "citations": [
            {
                "title": row.title,
                "url": row.canonical_url,
                "source_name": row.source_name,
                "quoted_text": row.quoted_text,
            }
            for row in response.citations.filter(url_status="safe")
        ],
    }


def report_history(*, user, subject_id) -> list[dict[str, Any]]:
    return [
        report_payload(row) for row in GeoReport.objects.filter(user=user, subject_id=subject_id)
    ]


def report_trends(*, user, subject_id) -> list[dict[str, Any]]:
    previous_by_signature: dict[tuple[str, str, str], GeoReport] = {}
    items = []
    rows = GeoReport.objects.filter(user=user, subject_id=subject_id).order_by("generated_at", "id")
    for report in rows:
        signature = (
            report.question_signature,
            report.model_signature,
            report.scoring_rule_version,
        )
        baseline = previous_by_signature.get(signature)
        items.append(
            {
                "report_id": str(report.pk),
                "generated_at": report.generated_at,
                "subject_version_id": str(report.subject_version_id),
                "geo_score": report.summary["geo"]["score"],
                "comparison": comparison(report, baseline) if baseline is not None else None,
            }
        )
        previous_by_signature[signature] = report
    return items


def _export_bytes(report: GeoReport, format: str) -> tuple[bytes, str]:
    title = f"GEO Report {report.pk}"
    export_payload: dict[str, Any] = {
        "report": report_payload(report),
        "questions": [],
    }
    total_questions = report.job.snapshot.questions.count()
    for page in range(1, max(1, (total_questions + 9) // 10) + 1):
        page_payload = question_group_page(report=report, page=page)
        for question in page_payload["results"]:
            question_copy = {**question, "results": []}
            for result in question["results"]:
                answer = (
                    full_answer(report=report, call_id=result["call_id"])["answer"]
                    if result["answer_available"]
                    else None
                )
                question_copy["results"].append({**result, "full_answer": answer})
            export_payload["questions"].append(question_copy)
    export_payload = _json_value(export_payload)
    if format == "pdf":
        from reportlab.pdfbase import pdfmetrics  # type: ignore[import-untyped]
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont  # type: ignore[import-untyped]
        from reportlab.pdfgen import canvas  # type: ignore[import-untyped]

        output = io.BytesIO()
        pdf = canvas.Canvas(output)
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        pdf.setFont("STSong-Light", 16)
        pdf.drawString(72, 800, title)
        pdf.setFont("STSong-Light", 10)
        lines = json.dumps(export_payload, ensure_ascii=False, indent=2).splitlines()
        y = 772
        for line in lines:
            for wrapped in textwrap.wrap(line, width=78, replace_whitespace=False) or [""]:
                if y < 60:
                    pdf.showPage()
                    pdf.setFont("STSong-Light", 10)
                    y = 800
                pdf.drawString(54, y, wrapped)
                y -= 14
        pdf.save()
        return output.getvalue(), "application/pdf"
    if format == "word":
        from docx import Document

        document = Document()
        document.add_heading(title, 0)
        document.add_heading("报告摘要", level=1)
        document.add_paragraph(json.dumps(export_payload, ensure_ascii=False, indent=2))
        output = io.BytesIO()
        document.save(output)
        return (
            output.getvalue(),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    if format == "excel":
        from openpyxl import Workbook  # type: ignore[import-untyped]

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Summary"
        sheet.append(["metric", "value"])
        for key, value in report.summary.items():
            sheet.append([key, json.dumps(value, ensure_ascii=False)])
        provenance = workbook.create_sheet("Provenance")
        provenance.append(["field", "value"])
        for key, value in report.provenance.items():
            provenance.append([key, json.dumps(value, ensure_ascii=False)])
        details = workbook.create_sheet("Question Details")
        details.append(
            [
                "question",
                "question_type",
                "model_key",
                "status",
                "snippet",
                "full_answer",
                "score",
                "citations",
            ]
        )
        for question in export_payload["questions"]:
            for result in question["results"]:
                details.append(
                    [
                        question["text"],
                        question["question_type"],
                        result["model_key"],
                        result["status"],
                        result["snippet"],
                        result["full_answer"],
                        json.dumps(result["score"], ensure_ascii=False),
                        json.dumps(result["citations"], ensure_ascii=False),
                    ]
                )
        output = io.BytesIO()
        workbook.save(output)
        return (
            output.getvalue(),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    raise ValueError("unsupported export format")


def create_export(*, report: GeoReport, user, format: str) -> ReportExport:
    if user.pk != report.user_id or format not in ReportExport.Format.values:
        raise Http404
    return ReportExport.objects.create(
        report=report,
        user=user,
        format=format,
        brand_snapshot={"brand_name": "显问 GEO", "white_label": False},
    )


def execute_export(*, export_id) -> ReportExport:
    export = ReportExport.objects.select_related("report").get(pk=export_id)
    export.status = ReportExport.Status.RUNNING
    export.save(update_fields=("status",))
    try:
        data, content_type = _export_bytes(export.report, export.format)
        extensions = {"pdf": "pdf", "word": "docx", "excel": "xlsx"}
        key = f"system/report-exports/{export.report_id}/{export.pk}.{extensions[export.format]}"
        storage_provider().put_system_object(key=key, data=data, content_type=content_type)
        export.status = ReportExport.Status.SUCCEEDED
        export.object_key = key
        export.finished_at = timezone.now()
        export.expires_at = timezone.now() + timedelta(days=7)
        export.save(update_fields=("status", "object_key", "finished_at", "expires_at"))
    except Exception:
        export.status = ReportExport.Status.FAILED
        export.safe_error_code = "REPORT_EXPORT_FAILED"
        export.finished_at = timezone.now()
        export.save(update_fields=("status", "safe_error_code", "finished_at"))
    return export
